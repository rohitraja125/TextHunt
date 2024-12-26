import os
import fitz
from collections import defaultdict
from preprocessing import preprocess_query, preprocess_text
from ocr import ocr_image, highlight_text_in_image
from nltk.stem import PorterStemmer
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from PIL import Image, ImageDraw
from spire.doc import Document

# new imports
from docx2pdf import convert
import pythoncom

def convert_to_json_output(results):
    result = []

    for file_paths, numbers in results.items():
        result.append([file_paths, list(numbers)])
    return result


def convert_docx_to_pdf(docx_path, pdf_path):
   
    
    # Initialize COM for the current thread
    pythoncom.CoInitialize()
    try:
        convert(docx_path, pdf_path)
    finally:
        # Uninitialize COM after the operation
        pythoncom.CoUninitialize()
    
    return pdf_path

def get_results(query, input_documents, isWordSearch):
    if isWordSearch.lower() == 'true':
        tokens = query.split()
        query = preprocess_query(query)
        query = preprocess_text(query, PorterStemmer())
        tokens = set(tokens)
        ranked_documents = get_relevant_results(query, input_documents)
        results = get_relevant_docs_with_pageno(ranked_documents, tokens)
        return results
    else:
        results = get_phrase_match(query, input_documents)
        return results

def get_relevant_docs_with_pageno(ranked_documents, tokens):
    results = defaultdict(set)
    for similarity, doc_path in ranked_documents:
        if similarity > 0:
            if is_pdf(doc_path):
                pdf_document = fitz.open(doc_path)
                for page_num in range(len(pdf_document)):
                    page = pdf_document[page_num]
                    words = page.get_text("words")
                    for word in words:
                        word_text = word[4]
                        for token in tokens:
                            if token.lower() in word_text.lower():
                                results[get_highlighted_file_output_path(doc_path)].add(page_num + 1)
                                text_instances = page.search_for(word_text)
                                for instance in text_instances:
                                    page.add_highlight_annot(instance)
                                break
                pdf_document.save(get_highlighted_file_output_path(doc_path))
                pdf_document.close()
            elif is_image(doc_path):
                highlight_text_in_image(tokens, doc_path, results)
            elif is_docx(doc_path):
                highlight_text_in_docx(doc_path, tokens, results)    
    return results

def get_highlighted_file_output_path(doc_path):
    return "highlighted_" + doc_path.split("/")[-1]

def is_pdf(document):
    return document.split(".")[-1] == "pdf"

def is_image(document):
    return document.split(".")[-1] in ('jpg', 'png', 'jpeg')

def is_docx(document):
    return document.split(".")[-1] == "docx"
    
def get_text_from_pdf(pdf_path):
    pdf_document = fitz.open(pdf_path)
    text = ""
    for page_num in range(len(pdf_document)):
        page = pdf_document[page_num]
        text += page.get_text()
    return text

def get_phrase_match(query, input_documents):
    relevant_docs = defaultdict(set)
    for i in range(len(input_documents)):
        is_matched = False
        if is_pdf(input_documents[i]):
            pdf_document = fitz.open(input_documents[i])
            for page_num in range(len(pdf_document)):
                page = pdf_document[page_num]
                text_instances = page.search_for(query)
                if len(text_instances) > 0:
                    is_matched = True
                    relevant_docs[get_highlighted_file_output_path(input_documents[i])].add(page_num + 1)
                for instance in text_instances:
                    page.add_highlight_annot(instance)
            if is_matched:
                pdf_document.save(get_highlighted_file_output_path(input_documents[i]))
            pdf_document.close()
        elif is_image(input_documents[i]):
            text = ocr_image(input_documents[i])
            if query.lower() in text.lower():
                relevant_docs[get_highlighted_file_output_path(input_documents[i])].add(1)
                image = Image.open(input_documents[i])
                draw = ImageDraw.Draw(image)
                image.save(get_highlighted_file_output_path(input_documents[i]))
        elif is_docx(input_documents[i]):
            highlight_phrase_in_docx(input_documents[i], query, relevant_docs)
    return relevant_docs

def get_relevant_results(query, input_documents):
    docs = []
    porter_stemmer = PorterStemmer()

    for i in range(len(input_documents)):
        if is_pdf(input_documents[i]):
            text = get_text_from_pdf(input_documents[i])
            stemmed_text = preprocess_text(text, porter_stemmer)
            docs.append(stemmed_text)
        elif is_image(input_documents[i]):
            text = ocr_image(input_documents[i])
            stemmed_text = preprocess_text(text, porter_stemmer)
            docs.append(stemmed_text)
        elif is_docx(input_documents[i]):
            text = get_text_from_docx(input_documents[i])
            stemmed_text = preprocess_text(text, porter_stemmer)
            # print("docx path is:", input_documents[i])
            pdf_path = input_documents[i].split("/")[-1].split(".")[0] + ".pdf";
            print("path: ", pdf_path)
            convert_docx_to_pdf(input_documents[i], pdf_path);
            docs.append(stemmed_text)

    results = proximity_search(query, docs, input_documents)
    return results

def proximity_search(query, documents, doc_paths):
    tfidf_vectorizer = TfidfVectorizer(stop_words='english')
    tfidf_matrix = tfidf_vectorizer.fit_transform(documents)
    query_vector = tfidf_vectorizer.transform([query])
    similarities = cosine_similarity(query_vector, tfidf_matrix)
    ranked_docs = sorted(zip(similarities[0], doc_paths), reverse=True)
    return ranked_docs


from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def get_text_from_docx(docx_file):
    doc = Document(docx_file)
    print(doc)
    text = []
    for paragraph in doc.paragraphs:
        text.append(paragraph.text)
    
    return "\n".join(text)

from spire.doc import Document as SpireDocument
from spire.doc.common import *

def highlight_text_in_docx(doc_path, tokens, results):
    document = SpireDocument()
    document.LoadFromFile(doc_path)
    isMatch = False
    for token in tokens:
        textSelections = document.FindAllString(token, False, False)    
        # Loop through all the instances
        for selection in textSelections:
            isMatch = True
        # Get the current instance as a single text range
            textRange = selection.GetAsOneRange()
            # Highlight the text range with a color
            textRange.CharacterFormat.HighlightColor = Color.get_Yellow()
    # Save the resulting document
    if isMatch:
        results[get_highlighted_file_output_path(doc_path=doc_path)].add(1)
        document.SaveToFile(get_highlighted_file_output_path(doc_path))
    document.Close()


def highlight_phrase_in_docx(doc_path, query, relevant_docs):
    document = SpireDocument()
    document.LoadFromFile(doc_path)
    textSelections = document.FindAllString(query, False, True)    
    for selection in textSelections:
        # Get the current instance as a single text range
        textRange = selection.GetAsOneRange()
        # Highlight the text range with a color
        textRange.CharacterFormat.HighlightColor = Color.get_Yellow()    
    # Save the resulting document
    if(len(textSelections) > 0):
        relevant_docs[get_highlighted_file_output_path(doc_path)].add(1)
        document.SaveToFile(get_highlighted_file_output_path(doc_path))
    document.Close()

